"""Export endpoints for generating DOCX and PDF documents."""

import io
from enum import Enum

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

router = APIRouter(prefix="/export", tags=["Export"])


class ExportFormat(str, Enum):
    DOCX = "docx"
    PDF = "pdf"


class ExportSourceReference(BaseModel):
    document_title: str
    chunk_index: int | None = None
    relevance_score: float | None = None


class ExportSection(BaseModel):
    heading: str
    content: str
    sources: list[ExportSourceReference] = []


class ExportRequest(BaseModel):
    sections: list[ExportSection]
    format: ExportFormat
    document_title: str = "Generated Document"


def _export_docx(req: ExportRequest) -> io.BytesIO:
    """Generate a DOCX file from export sections."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Title
    title = doc.add_heading(req.document_title, level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(13, 148, 136)

    # Sections
    all_sources: list[ExportSourceReference] = []
    for section in req.sections:
        doc.add_heading(section.heading, level=1)
        for paragraph_text in section.content.split("\n"):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip())
        all_sources.extend(section.sources)

    # Citations appendix
    if all_sources:
        doc.add_page_break()
        doc.add_heading("Sources", level=1)
        seen = set()
        for i, source in enumerate(all_sources, 1):
            key = source.document_title
            if key in seen:
                continue
            seen.add(key)
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(source.document_title)
            run.font.size = Pt(10)
            if source.relevance_score is not None:
                score_run = p.add_run(f" (relevance: {source.relevance_score:.0%})")
                score_run.font.size = Pt(9)
                score_run.font.color.rgb = RGBColor(128, 128, 128)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _export_pdf(req: ExportRequest) -> io.BytesIO:
    """Generate a PDF file from export sections."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(13, 148, 136)
    pdf.cell(0, 15, req.document_title, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # Sections
    all_sources: list[ExportSourceReference] = []
    for section in req.sections:
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(30, 30, 30)
        pdf.cell(0, 10, section.heading, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(50, 50, 50)
        # Write content, handling encoding
        content = section.content.encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, 6, content)
        pdf.ln(4)
        all_sources.extend(section.sources)

    # Citations appendix
    if all_sources:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(30, 30, 30)
        pdf.cell(0, 10, "Sources", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

        seen = set()
        idx = 1
        for source in all_sources:
            key = source.document_title
            if key in seen:
                continue
            seen.add(key)
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(50, 50, 50)
            line = f"{idx}. {source.document_title}"
            if source.relevance_score is not None:
                line += f" (relevance: {source.relevance_score:.0%})"
            line = line.encode("latin-1", "replace").decode("latin-1")
            pdf.cell(0, 7, line, new_x="LMARGIN", new_y="NEXT")
            idx += 1

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


@router.post("")
async def export_document(req: ExportRequest) -> StreamingResponse:
    """Export generated content as DOCX or PDF."""
    try:
        if req.format == ExportFormat.DOCX:
            buf = _export_docx(req)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ext = "docx"
        else:
            buf = _export_pdf(req)
            media_type = "application/pdf"
            ext = "pdf"

        filename = req.document_title.replace(" ", "_") + f".{ext}"
        return StreamingResponse(
            buf,
            media_type=media_type,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {e}") from e
